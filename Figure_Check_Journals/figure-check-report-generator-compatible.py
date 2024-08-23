import os
import sys
from PIL import Image
import fitz  # PyMuPDF
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def check_image(file_path):
    results = {
        "File": os.path.basename(file_path),
        "Type": "Image",
        "RGB Mode": "✓",
        "Resolution": "✓",
        "Dimensions": "✓",
        "Font": "N/A",
        "Font Size": "N/A"
    }
    try:
        with Image.open(file_path) as img:
            if img.mode != 'RGB':
                results["RGB Mode"] = f"✗ ({img.mode})"
            
            dpi = img.info.get('dpi')
            if dpi:
                if dpi[0] < 300 or dpi[1] < 300:
                    results["Resolution"] = f"✗ ({dpi[0]}x{dpi[1]} DPI)"
            else:
                results["Resolution"] = "? (Unable to determine)"
            
            width, height = img.size
            width_mm = width * 25.4 / 300  # Convert pixels to mm at 300 DPI
            height_mm = height * 25.4 / 300
            if width_mm > 180 or height_mm > 180:
                results["Dimensions"] = f"✗ ({width_mm:.1f}x{height_mm:.1f}mm)"
    except Exception as e:
        results = {key: f"Error: {str(e)}" for key in results}
    return results

def check_pdf(file_path):
    results = {
        "File": os.path.basename(file_path),
        "Type": "PDF",
        "RGB Mode": "? (Unable to determine)",
        "Resolution": "✓",
        "Dimensions": "✓",
        "Font": "✓",
        "Font Size": "✓"
    }
    try:
        doc = fitz.open(file_path)
        
        # Check if PDF contains raster images and their resolution
        for page in doc:
            for img in page.get_images():
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                width_mm = pix.width * 25.4 / 300  # Convert pixels to mm at 300 DPI
                height_mm = pix.height * 25.4 / 300
                if width_mm > 180 or height_mm > 180:
                    results["Dimensions"] = f"✗ ({width_mm:.1f}x{height_mm:.1f}mm)"
                if pix.xres < 300 or pix.yres < 300:
                    results["Resolution"] = f"✗ ({pix.xres}x{pix.yres} DPI)"
                break  # Check only the first image
        
        # Check for text properties
        for page in doc:
            text = page.get_text("dict")
            for block in text["blocks"]:
                if block["type"] == 0:  # Text block
                    for line in block["lines"]:
                        for span in line["spans"]:
                            font = span["font"]
                            if not re.search(r'helvetica|arial', font, re.IGNORECASE):
                                results["Font"] = f"✗ ({font})"
                            size = span["size"]
                            if size > 7 or size < 5:
                                results["Font Size"] = f"✗ ({size:.1f}pt)"
    except Exception as e:
        results = {key: f"Error: {str(e)}" for key in results}
    return results

def generate_report(results, output_file):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Figure Check Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary paragraph
    summary = doc.add_paragraph('This report summarizes the results of the figure checks based on Nature Genetics submission guidelines. Each figure has been checked for resolution, dimensions, font type, and font size. RGB color mode could not be automatically determined for PDFs.')
    
    # Add table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(["File", "Type", "RGB Mode", "Resolution", "Dimensions", "Font", "Font Size"]):
        hdr_cells[i].text = header
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
    
    for result in results:
        row_cells = table.add_row().cells
        for i, key in enumerate(["File", "Type", "RGB Mode", "Resolution", "Dimensions", "Font", "Font Size"]):
            cell = row_cells[i]
            cell.text = result[key]
            if '✗' in result[key]:
                run = cell.paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for issues
    
    # Add explanatory notes
    doc.add_heading('Notes:', level=2)
    notes = [
        "✓ indicates the criterion is met.",
        "✗ indicates an issue that needs to be addressed.",
        "? indicates the information could not be determined.",
        "N/A indicates the criterion is not applicable to this file type.",
        "RGB Mode: All figures should be in RGB color mode. For PDFs, please check this manually.",
        "Resolution: All figures should have a minimum resolution of 300 DPI.",
        "Dimensions: Maximum figure width is 180 mm for two-column figures.",
        "Font: Use sans-serif typeface, preferably Helvetica or Arial.",
        "Font Size: Text size should be between 5-7 points."
    ]
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    doc.save(output_file)

def main(directory):
    results = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
            results.append(check_image(file_path))
        elif filename.lower().endswith('.pdf'):
            results.append(check_pdf(file_path))
    
    output_file = os.path.join(directory, 'Figure_Check_Report.docx')
    generate_report(results, output_file)
    print(f"Report generated: {output_file}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python script.py <directory>")
    else:
        main(sys.argv[1])
